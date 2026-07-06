"""
참고 자료 파서 + 워크플로
- PDF, DOCX, XLSX, TXT 등 다형식 텍스트/표 추출
- 참고자료 요약 생성
- 요약을 작업 HWP/HWPX 문서 끝에 삽입
"""

from __future__ import annotations

import io
import os
import re
import time
from dataclasses import dataclass, field
from typing import Optional


@dataclass
class ReferenceDocument:
    filename: str
    file_type: str
    full_text: str = ''
    paragraphs: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)


def parse_reference(file_bytes: bytes, filename: str) -> ReferenceDocument:
    ext = os.path.splitext(filename)[1].lower()
    parsers = {
        '.txt': _parse_txt,
        '.md': _parse_txt,
        '.csv': _parse_txt,
        '.docx': _parse_docx,
        '.xlsx': _parse_xlsx,
        '.xls': _parse_xlsx,
        '.pdf': _parse_pdf,
        '.hwp': _parse_hwp_ref,
        '.hwpx': _parse_hwpx_ref,
    }
    parser = parsers.get(ext)
    if parser is None:
        doc = ReferenceDocument(filename=filename, file_type=ext)
        doc.errors.append(f'지원하지 않는 참고 자료 형식: {ext}')
        return doc
    return parser(file_bytes, filename)


def _parse_txt(file_bytes: bytes, filename: str) -> ReferenceDocument:
    doc = ReferenceDocument(filename=filename, file_type='txt')
    try:
        text = file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        text = file_bytes.decode('cp949', errors='ignore')
    doc.full_text = text
    doc.paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    return doc


def _parse_docx(file_bytes: bytes, filename: str) -> ReferenceDocument:
    doc = ReferenceDocument(filename=filename, file_type='docx')
    try:
        from docx import Document
    except ImportError:
        doc.errors.append('python-docx 미설치: pip install python-docx')
        return doc
    try:
        document = Document(io.BytesIO(file_bytes))
        for para in document.paragraphs:
            t = para.text.strip()
            if t:
                doc.paragraphs.append(t)
        for table in document.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                doc.tables.append(rows)
        doc.full_text = '\n'.join(doc.paragraphs)
        for i, tbl in enumerate(doc.tables):
            doc.full_text += f'\n\n[표 {i+1}]\n'
            for row in tbl:
                doc.full_text += ' | '.join(row) + '\n'
    except Exception as e:
        doc.errors.append(f'DOCX 파싱 오류: {e}')
    return doc


def _parse_xlsx(file_bytes: bytes, filename: str) -> ReferenceDocument:
    doc = ReferenceDocument(filename=filename, file_type='xlsx')
    try:
        from openpyxl import load_workbook
    except ImportError:
        doc.errors.append('openpyxl 미설치')
        return doc
    try:
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else '' for c in row]
                if any(c.strip() for c in cells):
                    rows.append(cells)
            if rows:
                doc.tables.append(rows)
                parts.append(f'[시트: {sheet_name}]')
                for r in rows[:50]:
                    parts.append(' | '.join(r))
        doc.full_text = '\n'.join(parts)
        doc.paragraphs = parts[:20]
        wb.close()
    except Exception as e:
        doc.errors.append(f'XLSX 파싱 오류: {e}')
    return doc


def _parse_pdf(file_bytes: bytes, filename: str) -> ReferenceDocument:
    doc = ReferenceDocument(filename=filename, file_type='pdf')
    try:
        from pypdf import PdfReader
    except ImportError:
        doc.errors.append('pypdf 미설치: pip install pypdf')
        return doc
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []
        for page in reader.pages:
            t = page.extract_text() or ''
            if t.strip():
                parts.append(t.strip())
        doc.full_text = '\n\n'.join(parts)
        doc.paragraphs = [p.strip() for p in doc.full_text.split('\n') if p.strip()]
    except Exception as e:
        doc.errors.append(f'PDF 파싱 오류: {e}')
    return doc


def _parse_hwp_ref(file_bytes: bytes, filename: str) -> ReferenceDocument:
    from hwp_core.hwp_parser import parse_document
    parsed = parse_document(file_bytes=file_bytes, filename=filename)
    doc = ReferenceDocument(filename=filename, file_type='hwp')
    doc.full_text = parsed.full_text
    doc.paragraphs = list(parsed.paragraphs)
    doc.errors = list(parsed.errors)
    for tbl in parsed.tables_raw:
        if tbl.get('rows'):
            doc.tables.append(tbl['rows'])
    return doc


def _parse_hwpx_ref(file_bytes: bytes, filename: str) -> ReferenceDocument:
    return _parse_hwp_ref(file_bytes, filename)


def build_reference_context(refs: list[ReferenceDocument], max_chars: int = 12000) -> str:
    """여러 참고 자료를 LLM 컨텍스트 문자열로 합침."""
    parts = []
    total = 0
    for ref in refs:
        header = f'## 참고자료: {ref.filename} ({ref.file_type})'
        body = ref.full_text[:max_chars - total - len(header) - 10]
        if not body.strip():
            continue
        chunk = f'{header}\n{body}'
        parts.append(chunk)
        total += len(chunk)
        if total >= max_chars:
            break
    return '\n\n'.join(parts)


# =========================================================
# 참고자료 워크플로 (구 reference_workflow.py)
# =========================================================

def normalize_insert_body(text: str) -> str:
    """마크다운·UI 힌트를 한글 문서 삽입용 평문으로 변환."""
    t = (text or '').strip()
    t = re.sub(r'👉[^\n]*', '', t)
    lines: list[str] = []
    for raw in t.splitlines():
        line = raw.strip()
        if not line or line in ('---', '***', '___'):
            continue
        line = re.sub(r'^#{1,6}\s*', '', line)
        line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        line = re.sub(r'\*([^*]+)\*', r'\1', line)
        line = re.sub(r'^[-*•]\s+', '· ', line)
        if line.startswith('-'):
            line = '· ' + line.lstrip('-').strip()
        lines.append(line)
    return '\n'.join(lines)


def generate_reference_summary(
    reference_context: str,
    model: str,
    ollama_url: str,
    focus: str = '',
) -> tuple[str, str]:
    """참고자료 전체를 작업 문서에 넣기 좋은 요약으로 변환."""
    from hwp_core.llm_client import generate

    if not reference_context.strip():
        return '', '참고자료가 없습니다. 사이드바에서 PDF·DOCX 등을 업로드하세요.'

    focus_line = f'\n특히 다음 관점을 반영하세요: {focus}' if focus else ''
    prompt = f"""다음 참고자료를 한글 공문서에 붙여 넣을 수 있는 요약으로 작성하세요.

규칙:
- 마크다운(---, ###, **) 사용 금지. 일반 문장과 · 불릿만 사용
- 제목 한 줄 후 본문 (번호 목록 가능)
- 표·수치는 빠짐없이 포함
- 800~2000자 내외{focus_line}

[참고자료]
{reference_context[:10000]}
"""
    result = generate(
        prompt, model, ollama_url,
        temperature=0.3, num_predict=3000, num_ctx=16384, timeout=180,
    )
    if result.get('error'):
        return '', f"LLM 오류: {result['error']}"
    text = result.get('text', '').strip()
    return normalize_insert_body(text), ''


def pick_summary_text(
    command: str,
    chat_history: list | None,
    cached_summary: str = '',
) -> str:
    """명령·캐시·이전 채팅에서 삽입할 본문 선택."""
    if cached_summary and len(cached_summary.strip()) >= 40:
        return normalize_insert_body(cached_summary)

    if chat_history:
        for msg in reversed(chat_history):
            if msg.get('role') != 'assistant':
                continue
            content = (msg.get('content') or '').strip()
            if len(content) < 40:
                continue
            if content.startswith('\U0001f449') or '모두 적용' in content:
                continue
            if '참고자료' in content or '요약' in content or len(content) > 200:
                return normalize_insert_body(content)

    colon = re.search(r'[:：]\s*(.+)$', command, re.S)
    if colon and len(colon.group(1).strip()) >= 40:
        return normalize_insert_body(colon.group(1))

    return ''


def propose_append_at_end(editor, body: str):
    """문서 마지막 문단 뒤에 본문 삽입 제안 (HWPX XML, hwpilot 불필요)."""
    from hwp_core.hwpx_editor import PendingChange

    body = normalize_insert_body(body)
    if not body or len(body) < 10:
        return None
    paras = editor.get_paragraphs()
    if not paras:
        return None
    last = paras[-1]
    import uuid
    change = PendingChange(
        id=str(uuid.uuid4())[:8],
        change_type='insert_after',
        location='문서 끝',
        old_text='',
        new_text=body,
        paragraph_index=last['index'],
        section_file=last.get('section_file'),
        search_hint='__END__',
    )
    editor.pending_changes.append(change)
    editor._bump_preview()
    return change


def append_summary_to_document(
    editor,
    command: str,
    reference_context: str,
    model: str,
    ollama_url: str,
    chat_history: list | None = None,
    cached_summary: str = '',
    source_filename: str = 'doc.hwpx',
    file_bytes: bytes | None = None,
) -> dict:
    """참고자료 요약을 작업 문서 끝에 추가."""
    start = time.time()
    body = pick_summary_text(command, chat_history, cached_summary)

    if not body and reference_context:
        body, err = generate_reference_summary(reference_context, model, ollama_url)
        if err:
            return {
                'type': 'edit', 'intent': 'append_ref',
                'message': f'요약 생성 실패: {err}',
                'changes': 0, 'elapsed': round(time.time() - start, 1),
            }

    if not body:
        return {
            'type': 'edit', 'intent': 'append_ref',
            'message': (
                '삽입할 요약이 없습니다. 먼저 「참고자료 요약 생성」을 누르거나 '
                '「참고자료 내용 알려줘」로 요약을 받은 뒤 「작업 문서 끝에 추가」를 사용하세요.'
            ),
            'changes': 0, 'elapsed': round(time.time() - start, 1),
        }

    title = '【참고자료 요약】'
    full_body = f'{title}\n{body}'

    if editor is not None:
        ch = propose_append_at_end(editor, full_body)
        if ch:
            return {
                'type': 'edit', 'intent': 'append_ref',
                'message': (
                    f'참고자료 요약 {len(body)}자를 문서 끝에 제안했습니다. '
                    '왼쪽 **노란색** 확인 후 「모두 적용」하세요.'
                ),
                'changes': 1,
                'elapsed': round(time.time() - start, 1),
                'summary_text': body,
            }

    from hwp_core.hwp_backends import get_backend_status, hwpilot_apply_content
    if file_bytes and get_backend_status().hwpilot:
        new_bytes, msg = hwpilot_apply_content(
            file_bytes, source_filename, full_body, anchor='__END__',
        )
        if new_bytes:
            return {
                'type': 'edit', 'intent': 'append_ref',
                'message': f'{msg} — 참고자료 요약이 문서 끝에 반영되었습니다.',
                'changes': 1,
                'elapsed': round(time.time() - start, 1),
                'applied_direct': True,
                'new_file_bytes': new_bytes,
                'summary_text': body,
            }
        return {
            'type': 'edit', 'intent': 'append_ref',
            'message': msg or '문서 끝 삽입에 실패했습니다.',
            'changes': 0, 'elapsed': round(time.time() - start, 1),
        }

    return {
        'type': 'edit', 'intent': 'append_ref',
        'message': '작업 문서가 열려 있지 않습니다. HWP/HWPX를 업로드하세요.',
        'changes': 0, 'elapsed': round(time.time() - start, 1),
    }
