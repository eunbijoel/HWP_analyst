"""
참고 자료 파서 — PDF, DOCX, XLSX, TXT 등 다형식 텍스트/표 추출
AI 편집·Q&A 컨텍스트용
"""

import io
import os
import re
from dataclasses import dataclass, field
from typing import Optional


@dataclass
class ReferenceDocument:
    filename: str
    file_type: str
    full_text: str = ''
    paragraphs: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)


def parse_reference(file_bytes: bytes, filename: str) -> ReferenceDocument:
    ext = os.path.splitext(filename)[1].lower()
    parsers = {
        '.txt': _parse_txt,
        '.md': _parse_txt,
        '.csv': _parse_txt,
        '.docx': _parse_docx,
        '.xlsx': _parse_xlsx,
        '.xls': _parse_xlsx,
        '.pdf': _parse_pdf,
        '.hwp': _parse_hwp_ref,
        '.hwpx': _parse_hwpx_ref,
    }
    parser = parsers.get(ext)
    if parser is None:
        doc = ReferenceDocument(filename=filename, file_type=ext)
        doc.errors.append(f'지원하지 않는 참고 자료 형식: {ext}')
        return doc
    return parser(file_bytes, filename)


def _parse_txt(file_bytes: bytes, filename: str) -> ReferenceDocument:
    doc = ReferenceDocument(filename=filename, file_type='txt')
    try:
        text = file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        text = file_bytes.decode('cp949', errors='ignore')
    doc.full_text = text
    doc.paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    return doc


def _parse_docx(file_bytes: bytes, filename: str) -> ReferenceDocument:
    doc = ReferenceDocument(filename=filename, file_type='docx')
    try:
        from docx import Document
    except ImportError:
        doc.errors.append('python-docx 미설치: pip install python-docx')
        return doc
    try:
        document = Document(io.BytesIO(file_bytes))
        for para in document.paragraphs:
            t = para.text.strip()
            if t:
                doc.paragraphs.append(t)
        for table in document.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                doc.tables.append(rows)
        doc.full_text = '\n'.join(doc.paragraphs)
        for i, tbl in enumerate(doc.tables):
            doc.full_text += f'\n\n[표 {i+1}]\n'
            for row in tbl:
                doc.full_text += ' | '.join(row) + '\n'
    except Exception as e:
        doc.errors.append(f'DOCX 파싱 오류: {e}')
    return doc


def _parse_xlsx(file_bytes: bytes, filename: str) -> ReferenceDocument:
    doc = ReferenceDocument(filename=filename, file_type='xlsx')
    try:
        from openpyxl import load_workbook
    except ImportError:
        doc.errors.append('openpyxl 미설치')
        return doc
    try:
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else '' for c in row]
                if any(c.strip() for c in cells):
                    rows.append(cells)
            if rows:
                doc.tables.append(rows)
                parts.append(f'[시트: {sheet_name}]')
                for r in rows[:50]:
                    parts.append(' | '.join(r))
        doc.full_text = '\n'.join(parts)
        doc.paragraphs = parts[:20]
        wb.close()
    except Exception as e:
        doc.errors.append(f'XLSX 파싱 오류: {e}')
    return doc


def _parse_pdf(file_bytes: bytes, filename: str) -> ReferenceDocument:
    doc = ReferenceDocument(filename=filename, file_type='pdf')
    try:
        from pypdf import PdfReader
    except ImportError:
        doc.errors.append('pypdf 미설치: pip install pypdf')
        return doc
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []
        for page in reader.pages:
            t = page.extract_text() or ''
            if t.strip():
                parts.append(t.strip())
        doc.full_text = '\n\n'.join(parts)
        doc.paragraphs = [p.strip() for p in doc.full_text.split('\n') if p.strip()]
    except Exception as e:
        doc.errors.append(f'PDF 파싱 오류: {e}')
    return doc


def _parse_hwp_ref(file_bytes: bytes, filename: str) -> ReferenceDocument:
    from hwp_parser import parse_document
    parsed = parse_document(file_bytes=file_bytes, filename=filename)
    doc = ReferenceDocument(filename=filename, file_type='hwp')
    doc.full_text = parsed.full_text
    doc.paragraphs = list(parsed.paragraphs)
    doc.errors = list(parsed.errors)
    for tbl in parsed.tables_raw:
        if tbl.get('rows'):
            doc.tables.append(tbl['rows'])
    return doc


def _parse_hwpx_ref(file_bytes: bytes, filename: str) -> ReferenceDocument:
    return _parse_hwp_ref(file_bytes, filename)


def build_reference_context(refs: list[ReferenceDocument], max_chars: int = 12000) -> str:
    """여러 참고 자료를 LLM 컨텍스트 문자열로 합침."""
    parts = []
    total = 0
    for ref in refs:
        header = f'## 참고자료: {ref.filename} ({ref.file_type})'
        body = ref.full_text[:max_chars - total - len(header) - 10]
        if not body.strip():
            continue
        chunk = f'{header}\n{body}'
        parts.append(chunk)
        total += len(chunk)
        if total >= max_chars:
            break
    return '\n\n'.join(parts)
